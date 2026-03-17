import os
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
os.environ["TRANSFORMERS_VERBOSITY"] = "error"

import pickle
import hashlib
import warnings
import glob as glob_module
from typing import List
from langchain_core.documents import Document
from langchain_community.document_loaders import DirectoryLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from rank_bm25 import BM25Okapi
from dotenv import load_dotenv
from src.config import (
    EMBED_MODEL, DB_DIR, CHUNKS_PKL, BM25_PKL, HASH_PKL,
    DATA_DIR, CHUNK_SIZE, CHUNK_OVERLAP,
)

warnings.filterwarnings("ignore")
load_dotenv()

_embeddings = None


def _get_embeddings():
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    return _embeddings


# ── File dedup (content hash) ─────────────────────────────────────────────────

def _file_hash(file_path: str) -> str:
    h = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _load_hashes() -> set:
    if os.path.exists(HASH_PKL):
        with open(HASH_PKL, "rb") as f:
            return pickle.load(f)
    return set()


def _save_hashes(hashes: set):
    with open(HASH_PKL, "wb") as f:
        pickle.dump(hashes, f)


def is_already_ingested(file_path: str) -> bool:
    return _file_hash(file_path) in _load_hashes()


# ── Multi-format loaders ──────────────────────────────────────────────────────

def _load_pdf(file_path: str) -> List[Document]:
    """PyMuPDF — handles text and image-heavy PDFs."""
    import fitz
    docs, pdf = [], fitz.open(file_path)
    for i, page in enumerate(pdf):
        text = page.get_text().strip()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={"source": file_path, "page": i + 1, "format": "pdf"},
            ))
    pdf.close()
    return docs


def _load_docx(file_path: str) -> List[Document]:
    """Extract paragraphs from Word documents."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not full_text.strip():
        return []
    return [Document(
        page_content=full_text,
        metadata={"source": file_path, "page": 1, "format": "docx"},
    )]


def _load_csv(file_path: str) -> List[Document]:
    """Load each CSV row as a document chunk."""
    import csv
    docs = []
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            text = " | ".join(f"{k}: {v}" for k, v in row.items() if v and v.strip())
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "page": i + 1, "format": "csv"},
                ))
    return docs


def _load_txt(file_path: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read().strip()
    if not content:
        return []
    return [Document(
        page_content=content,
        metadata={"source": file_path, "page": 1, "format": "txt"},
    )]


LOADERS = {
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".doc":  _load_docx,
    ".csv":  _load_csv,
    ".txt":  _load_txt,
    ".md":   _load_txt,
}

SUPPORTED_EXTENSIONS = list(LOADERS.keys())


def load_file(file_path: str) -> List[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    loader_fn = LOADERS.get(ext)
    if loader_fn is None:
        raise ValueError(f"Unsupported format: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    return loader_fn(file_path)


def load_documents(data_dir: str) -> List[Document]:
    docs = []
    for ext in SUPPORTED_EXTENSIONS:
        for path in glob_module.glob(os.path.join(data_dir, f"**/*{ext}"), recursive=True):
            try:
                docs.extend(load_file(path))
            except Exception as e:
                print(f"Warning: could not load {path}: {e}")
    return docs


# ── Chunking & indexing ───────────────────────────────────────────────────────

def chunk_documents(docs: List[Document], start_id: int = 0) -> List[Document]:
    docs = [d for d in docs if d.page_content and d.page_content.strip()]
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        add_start_index=True,
    )
    chunks = splitter.split_documents(docs)
    chunks = [c for c in chunks if c.page_content.strip()]
    for i, chunk in enumerate(chunks):
        chunk.metadata["doc_id"] = f"doc_{start_id + i}"
    return chunks


def build_dense_index(chunks: List[Document], persist_directory: str = DB_DIR):
    return Chroma.from_documents(
        documents=chunks,
        embedding=_get_embeddings(),
        persist_directory=persist_directory,
    )


def build_sparse_index(chunks: List[Document], save_path: str = BM25_PKL):
    bm25 = BM25Okapi([c.page_content.lower().split() for c in chunks])
    with open(save_path, "wb") as f:
        pickle.dump(bm25, f)
    return bm25


# ── Public API ────────────────────────────────────────────────────────────────

def ingest_uploaded_file(file_path: str) -> int:
    """Incrementally ingest one file. Returns number of new chunks added."""
    fhash  = _file_hash(file_path)
    hashes = _load_hashes()
    if fhash in hashes:
        return 0  # Already indexed

    docs = load_file(file_path)
    if not docs:
        raise ValueError(
            f"No text could be extracted from '{os.path.basename(file_path)}'. "
            "The file may be empty or a purely image-based scan."
        )

    existing_chunks = []
    if os.path.exists(CHUNKS_PKL):
        with open(CHUNKS_PKL, "rb") as f:
            existing_chunks = pickle.load(f)

    new_chunks = chunk_documents(docs, start_id=len(existing_chunks))
    if not new_chunks:
        raise ValueError("No valid text chunks could be created from this file.")

    all_chunks = existing_chunks + new_chunks

    # Update dense index
    vectorstore = Chroma(persist_directory=DB_DIR, embedding_function=_get_embeddings())
    vectorstore.add_documents(new_chunks)

    # Update sparse index & chunk store
    build_sparse_index(all_chunks)
    with open(CHUNKS_PKL, "wb") as f:
        pickle.dump(all_chunks, f)

    hashes.add(fhash)
    _save_hashes(hashes)

    # Reset cached RAG chain
    try:
        from src.generate import reset_chain
        reset_chain()
    except ImportError:
        pass

    return len(new_chunks)


def ingest_all(data_dir: str = DATA_DIR):
    """Full rebuild of both indexes from all files in data_dir."""
    os.makedirs(data_dir, exist_ok=True)
    docs = load_documents(data_dir)
    if not docs:
        print("No documents found.")
        return
    chunks = chunk_documents(docs)
    print(f"Chunked into {len(chunks)} segments.")
    with open(CHUNKS_PKL, "wb") as f:
        pickle.dump(chunks, f)
    build_dense_index(chunks)
    build_sparse_index(chunks)
    print("Ingestion complete.")


if __name__ == "__main__":
    os.makedirs(DATA_DIR, exist_ok=True)
    if not os.listdir(DATA_DIR):
        with open(os.path.join(DATA_DIR, "sample.txt"), "w", encoding="utf-8") as f:
            f.write(
                "LangChain is a framework for developing applications powered by LLMs.\n\n"
                "Cohere provides NLP models including an industry-leading reranker API.\n\n"
                "RAG (Retrieval-Augmented Generation) enhances LLMs with external knowledge."
            )
    ingest_all()
